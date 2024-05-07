from docx import Document
from docx.shared import Inches
import os
import re


class Loader:
    def __init__(self):
        self.global_index = 0

    def extract_content(self, docx_file):
        content_list = []
        current_chunk = []

        doc = Document(docx_file)

        for paragraph in doc.paragraphs:
            # if paragraph.text and paragraph.style.font.size >= Inches(0.2):
            if paragraph.text and paragraph.style.font.size is not None and paragraph.style.font.size >= Inches(0.2):

                if current_chunk:
                    content_list.append({"index": self.global_index, "content": "\n".join(current_chunk)})
                    current_chunk = []
                    self.global_index += 1

                current_chunk.append(paragraph.text)
            elif paragraph.text:
                current_chunk.append(paragraph.text)

        if current_chunk:
            content_list.append({"index": self.global_index, "content": "\n".join(current_chunk)})
            self.global_index += 1

        return content_list

    def get_images(self, docx_file):
        image_dict = {}
        doc = Document(docx_file)
        dict_rel = doc.part._rels
        image_index = 1

        for rel in dict_rel:
            rel = dict_rel[rel]
            if "image" in rel.target_ref:
                if not os.path.exists("./images"):
                    os.makedirs("./images")

                img_name = re.findall("/(.*)", rel.target_ref)[0]
                img_name = f'{self.global_index}-{image_index}.{img_name.split(".")[-1]}'

                with open(f'{"./images"}/{img_name}', "wb") as f:
                    f.write(rel.target_part.blob)

                image_dict[self.global_index] = image_dict.get(self.global_index, [])
                image_dict[self.global_index].append(img_name)
                image_index += 1

        return image_dict
