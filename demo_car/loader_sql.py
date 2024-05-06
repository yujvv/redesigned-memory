import docx
import sqlite3
from collections import defaultdict

def extract_content(doc):
    content = defaultdict(list)
    current_title = None

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("###"):
            current_title = paragraph.text.strip()[3:].strip()
        else:
            content[current_title].append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                row_content.append(cell.text.strip())
            content[current_title].append(row_content)

    for shape in doc.inline_shapes:
        content[current_title].append(shape)

    return content

def save_to_sqlite(content):
    conn = sqlite3.connect('content.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS content
                 (title TEXT, content TEXT, type TEXT)''')

    for title, items in content.items():
        for item in items:
            if isinstance(item, list):
                # Table row
                content_str = '|'.join(item)
                c.execute("INSERT INTO content VALUES (?, ?, 'table')", (title, content_str))
            elif isinstance(item, docx.shape.InlineShape):
                # Image
                c.execute("INSERT INTO content VALUES (?, ?, 'image')", (title, str(item)))
            else:
                # Text
                c.execute("INSERT INTO content VALUES (?, ?, 'text')", (title, item))

    conn.commit()
    conn.close()

def main():
    doc = docx.Document('2test.docx')
    content = extract_content(doc)

    print("字典的形状：", len(content), "个键值对")

    # 打印前五个键值对的内容
    print("前五个键值对的内容：")
    count = 0
    for key, value in content.items():
        if count < 5:
            print(key, ":", value)
            count += 1
        else:
            break


    # save_to_sqlite(content)

if __name__ == "__main__":
    main()