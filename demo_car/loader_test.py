from docx import Document
from docx.shared import Inches
from docx.oxml.shared import qn
import os
import re

# 定义全局变量用于保存当前的index
global_index = 0


def extract_content(docx_file):
    global global_index  # 声明使用全局变量

    doc = Document(docx_file)
    content_list = []
    # image_dict = {}
    current_chunk = []
    current_title = None

    for paragraph in doc.paragraphs:
        # if paragraph.text.startswith('###') and paragraph.style.font.size >= Inches(0.3):
        #     print("The 0.3 chunking____")            
        #     if current_chunk:
        #         # 使用全局index
        #         content_list.append({"index": global_index, "content": "\n".join(current_chunk)})
        #         current_chunk = []
        #         global_index += 1  # 每次添加chunk时递增全局index

        if paragraph.text and paragraph.style.font.size >= Inches(0.2):
            print("The 0.2 chunking____")
            if current_chunk:
                content_list.append({"index": global_index, "title": current_title, "content": "\n".join(current_chunk)})
                current_chunk = []
                global_index += 1
            current_title = paragraph.text

            current_chunk.append(paragraph.text)


        elif paragraph.text:
            current_chunk.append(paragraph.text)

    if current_chunk:
        print("The last chunking____")
        content_list.append({"index": global_index, "title": current_title, "content": "\n".join(current_chunk)})
        global_index += 1
        
    return content_list


def get_image(docx_file):
    image_dict = {}
    doc = Document(docx_file)
    dict_rel = doc.part._rels
    # dict_rel = paragraph.part._rels
    image_index = 1
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            if not os.path.exists("./images"):
                os.makedirs("./images")

            img_name = re.findall("/(.*)", rel.target_ref)[0]
            img_name = f'{global_index}-{image_index}.{img_name.split(".")[-1]}'

            with open(f'{"./images"}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)

            image_dict[global_index] = image_dict.get(global_index, [])
            image_dict[global_index].append(img_name)
            image_index += 1


if __name__ == '__main__':
    docx_file = '2test.docx'
    content_list = extract_content(docx_file)
    # image_dict = get_image(docx_file)

    print("Content List:")
    for i, chunk in enumerate(content_list):
        print(f"Chunk {i + 1}:\n{chunk}\n")
        print("----------")



    # print("Image Dictionary:")
    # print(image_dict)
